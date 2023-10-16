import hashlib
from array import array

import docx
from PIL import Image, ImageOps
from PyQt5.QtWidgets import QMessageBox
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK  # 오류나도 정상작동
from docx.shared import Inches, Pt, RGBColor
from datetime import datetime
import pandas as pd
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# 자동 문서 작업은 이 소스코드에서 할 것.

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#간격좁게, 왼쪽정렬
def p_spacing_left(paragraph_format):
    paragraph_format.line_spacing = Pt(19)
    paragraph_format.space_after = Pt(7)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def p_spacing_low_left(paragraph_format):
    paragraph_format.line_spacing = Pt(12)
    paragraph_format.space_after = Pt(1)
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#테이블 내용 들여쓰기
def indent_table(table, indent):
    # noinspection PyProtectedMember
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(indent))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)



def feat_importance(imgPath,document):
    """head = document.add_heading('결과 자료', level=1)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = head.paragraph_format
    paragraph_format.line_spacing = Pt(25)
    paragraph_format.space_after = Pt(5)"""

    feature = document.add_paragraph('[특성 중요도]')
    paragraph_format = feature.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_spacing_left(paragraph_format)
    # 'assets/feature_importance/DecisionTree_feature_importance Top20.png'
    feature = document.add_paragraph()
    paragraph_format = feature.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    im = Image.open(imgPath)

    im = im.crop((0, 30, 1750, 490))  # left, upper, right, lower

    # 이미지 테두리
    # im = ImageOps.expand(im, border=1, fill='black')
    # previously, image was 826 pixels wide, cropping to 825 pixels wide
    im.save(imgPath, dpi=(80, 80))
    plot_feature = feature.add_run()
    plot_feature.add_picture(imgPath, width=Inches(7.7), height=Inches(2.8))
    paragraph_format.left_indent = Pt(-65)

def integrityloganddicom(dicommodule, log, name1):
    if log.log.viewrexallog['Explaination'].str.contains(dicommodule.patient1.patient_id).any():
        matching_explanations = log.log.viewrexallog[log.log.viewrexallog['Explaination'].str.contains(dicommodule.patient1.patient_id)]['Time']
        for explanationtime in matching_explanations:
            print(explanationtime)
            ctime = dicommodule.patient1_ctime
            if ctime == explanationtime:
                print("해당파일은", explanationtime, "에 조작되었습니다")
            else :
                print("로그에 해당파일에 대한 수정행위는 있지만, 일치하는 시간은 없습니다.")
    else :
        print("There are no matching logs.")


#self.addressip, self.file1, self.file2, self.diffs,
#mod, name1, name2, list, dicommodule, logmodule
def make_docx(dicommodule, logmodule) -> None:
    name1 = dicommodule.file1
    name2 = dicommodule.file2
    list = dicommodule.diffs
    title = 'Medical Forgery Analysis Report'
    sub_title = name1
    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    dicom_filename = '- Filename : ' + name1
    score = '50'
    mod = 0
    # if count > 1: #리스트를 문자열로 변환하는 작업
    # predict = '\n'.join(predict)
    # else :
    # predict = ''.join(predict)


    document = Document()
    # document.add_picture('assets/logo.jpg', width=Inches(1.0)) #사진 정렬 명령어 찾지 못함

    head = document.add_heading(title, level=0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = head.paragraph_format
    paragraph_format.line_spacing = Pt(26)
    paragraph_format.space_after = Pt(10)



    time = document.add_paragraph('Report Create Time : ' + current_time )
    paragraph_format = time.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hash = document.add_paragraph('File hash : ' + dicommodule.dicom_filehash1)
    paragraph_format = hash.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    head2 = document.add_heading('DICOM', level=2)
    head2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = head2.paragraph_format
    paragraph_format.line_spacing = Pt(10)
    paragraph_format.space_after = Pt(1)

    #전체폰트
    """FStyle = paragraph.styles['Normal']
    FStyle.font.bold = True"""

    use_algorithm = document.add_paragraph('[File information]')
    paragraph_format = use_algorithm.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_spacing_left(paragraph_format)

    """    use_algorithm.style = document.styles['Normal']
        font = document.style.font
        font.name = 'Arial'
        font.size = Pt(10)
        font.bold = False"""

    use_algorithm1 = document.add_paragraph(dicom_filename)
    """use_algorithm1.style = document.styles['Normal']"""
    """use_algorithm1.style = FStyle"""


    paragraph_format = use_algorithm1.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_spacing_low_left(paragraph_format)


    explain_model = document.add_paragraph('')
    paragraph_format = explain_model.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT



    if mod == 0:
        explain_model.add_run('→ The DICOM is suspected of being tampered with.')
    elif mod == 1:
        explain_model.add_run('→ Suspicious activity was detected in the PACS.')

    paragraph_format.left_indent = Pt(5)
    p_spacing_low_left(paragraph_format)

    write_score = document.add_paragraph('[Accuracy]')
    paragraph_format = write_score.paragraph_format
    p_spacing_left(paragraph_format)

    write_score1 = document.add_paragraph(str(score))
    paragraph_format = write_score1.paragraph_format

    result = document.add_paragraph('[Result]')
    paragraph_format = result.paragraph_format
    p_spacing_left(paragraph_format)
    # paragraph_format.keep_with_next = True

    """if count < 5:
        result_table = document.add_table(rows=count + 1, cols=2)
        hdr_cells = result_table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Result'

        for i in range(1, count + 1):
            row = result_table.rows[i]
            row.cells[0].text = name_list[i - 1]
            row.cells[1].text = predict[i - 1]

    else:"""
    result_table = document.add_table(rows=3, cols=2)
    for cell in result_table.columns[0].cells:
        cell.width = Inches(8.0)

    for cell in result_table.columns[1].cells:
        cell.width = Inches(3.0)

    result_table.style = 'Table Grid'
    result_table.rows.style = "borderColor:red;background-color:gray"#'Table Grid'

    hdr_cells = result_table.rows[0].cells
    hdr_cells[0].text = 'Filename'
    p = hdr_cells[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].text = 'Forgery Position'
    p = hdr_cells[1].paragraphs[0]
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    firstfile = result_table.rows[1].cells
    secondfile = result_table.rows[2].cells

    firstfile[0].text = name1
    secondfile[0].text = name2
    firstfile[1].text = ','.join(list)
    #write_alert = result.add_run('\t전체 검사 결과는 Reports/csv_list/Results 디렉토리에 저장됩니다.')
    #write_alert.font.size = docx.shared.Pt(9)
    #write_alert.font.style = 'color:Gray'
    #write_alert.font.color.rgb = RGBColor(0xA6, 0xA6, 0xA6)

    #paragraph_format = write_alert.paragraph_format
    #paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    indent_table(result_table, '200')

    explain_model = document.add_paragraph('\n')
    paragraph_format = explain_model.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


    mapping = document.add_paragraph('[Combine Result]')
    paragraph_format = mapping.paragraph_format
    p_spacing_left(paragraph_format)
    integrityloganddicom(dicommodule, logmodule, name1)


    a = explain_model.add_run()
    a.add_break(WD_BREAK.PAGE)
    # result_add = result.add_run()

    # result.add_run('\nNote : KNN은 현재 시각화가 되지 않습니다.')
    # document.add_picture('assets/3.png',width=Inches(6.0))

    # a = document.add_paragraph('실험 1')  # add_run 사용 안할 시 텍스트 따로 그림이 따로 배치가 됨.(페이지 넘어갈때)
    # c = a.add_run()
    # c.add_picture('assets/4.png',width=Inches(6.0))

    # b = document.add_paragraph('실험 2')
    # document.add_picture('assets/2.png',width=Inches(6.0))
    if mod == 2:
        explain_model.add_run('악성코드의 판별에 기여한 특성의 중요도는 다음과 같습니다.')
    elif mod ==3:
        explain_model.add_run('악성코드의 판별에 기여한 특성의 중요도는 다음과 같습니다.')

    if mod == 2:
        imgPath = 'img/fake.png'
        feat_importance(imgPath, document)


    elif mod == 3:
        imgPath = 'img/fake.png'
        feat_importance(imgPath, document)

    if mod == 2:
        tree = document.add_paragraph('[트리 시각화]')
        #show_tree = tree.add_run()
        paragraph_format = tree.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_spacing_left(paragraph_format)

        tree1 = document.add_paragraph()
        paragraph_format = tree1.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


        imgPath = 'info.png'

        im = Image.open(imgPath)
        im = im.crop((200, 220, 1680, 1200)) #left, upper, right, lower
        im = ImageOps.expand(im, border=1, fill='black')
        # previously, image was 826 pixels wide, cropping to 825 pixels wide
        im.save(imgPath, dpi=(80, 80))
        plot_feature = tree1.add_run()
        plot_feature.add_picture(imgPath,width=Inches(6.4), height=Inches(4.8))
        paragraph_format.left_indent = Pt(5)


    head2 = document.add_heading('PACS', level=2)
    head2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = head2.paragraph_format
    paragraph_format.line_spacing = Pt(10)
    paragraph_format.space_after = Pt(1)



    alert = document.add_paragraph('\n\n\nCaution: This system cannot directly modify the data of the server, DICOM, or PACS.\n')
    paragraph_format = alert.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER




    reportPath = sub_title + '.docx'
    document.save(reportPath)

    msg = QMessageBox()
    msg.setWindowTitle("Complete")
    msg.setText("The report has been generated.")
    msg.setStyleSheet('font: 25 11pt "KoPubWorldDotum";background-color: rgb(255, 255, 255);')

    msg.exec_()

    return reportPath
